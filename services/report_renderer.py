import json
import os
import re
from datetime import datetime
from typing import Any, Dict, List, Tuple

from flask import current_app
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from models import Report, SATReport

TABLE_SECTION_KEYS: List[Tuple[str, str]] = [
    ('PRE_TEST_REQUIREMENTS', 'Pre-Test Requirements'),
    ('KEY_COMPONENTS', 'Key Components'),
    ('IP_RECORDS', 'IP Address Records'),
    ('RELATED_DOCUMENTS', 'Related Documents'),
    ('PRE_APPROVALS', 'Pre-Approvals'),
    ('POST_APPROVALS', 'Post-Approvals'),
    ('SIGNAL_LISTS', 'Digital Signal Checks'),
    ('ANALOGUE_LISTS', 'Analogue Signal Checks'),
    ('MODBUS_DIGITAL_LISTS', 'Modbus Digital Signals'),
    ('MODBUS_ANALOGUE_LISTS', 'Modbus Analogue Signals'),
    ('DATA_VALIDATION', 'Data Validation'),
    ('PROCESS_TEST', 'Process Test Matrix'),
    ('SCADA_VERIFICATION', 'SCADA Verification'),
    ('TRENDS_TESTING', 'Trends Testing'),
    ('ALARM_LIST', 'Alarm List'),
]


def generate_modern_sat_report(submission_id: str) -> Dict[str, Any]:
    report = Report.query.filter_by(id=submission_id).first()
    if not report:
        return {'error': 'Report not found.'}

    sat_report = SATReport.query.filter_by(report_id=submission_id).first()
    if not sat_report:
        return {'error': 'SAT data not available for this report.'}

    try:
        payload = json.loads(sat_report.data_json) if sat_report.data_json else {}
    except json.JSONDecodeError:
        current_app.logger.warning('Invalid SAT data JSON for %s', submission_id)
        payload = {}

    context = payload.get('context', payload) or {}

    doc = _create_document()
    _clear_document_body(doc)

    _build_cover_page(doc, report, context)
    doc.add_page_break()
    _add_summary_sections(doc, context)
    _add_table_sections(doc, context)

    output_dir = current_app.config.get('OUTPUT_DIR', 'outputs')
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f'SAT_Report_{submission_id}_Modern.docx')

    try:
        doc.save(output_path)
    except Exception as exc:  # noqa: BLE001 - propagate friendly error
        current_app.logger.error('Failed to save modern SAT report for %s: %s', submission_id, exc, exc_info=True)
        return {'error': 'Could not write the generated document to disk.'}

    return {
        'path': output_path,
        'download_name': _build_download_name(report, context, submission_id),
    }


def _create_document() -> Document:
    candidates: List[str] = []
    for key in ('MODERN_TEMPLATE_PATH', 'TEMPLATE_FILE'):
        candidate = current_app.config.get(key)
        if candidate and candidate not in candidates:
            candidates.append(candidate)
    candidates.append(os.path.join(current_app.root_path, 'templates', 'SAT_Template.docx'))

    for candidate in candidates:
        if candidate and os.path.exists(candidate):
            try:
                return Document(candidate)
            except Exception as exc:  # noqa: BLE001 - log and fallback
                current_app.logger.warning('Unable to load template %s: %s', candidate, exc)
    return Document()


def _clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for element in list(body):
        body.remove(element)


def _build_cover_page(doc: Document, report: Report, context: Dict[str, Any]) -> None:
    title = _clean_text(context.get('DOCUMENT_TITLE') or report.document_title or 'SAT Report')

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.name = 'Calibri'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Site Acceptance Test Report')
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.name = 'Calibri'

    doc.add_paragraph('')

    meta_rows = [
        ('Client', context.get('CLIENT_NAME') or report.client_name),
        ('Project Reference', context.get('PROJECT_REFERENCE') or report.project_reference),
        ('Document Reference', context.get('DOCUMENT_REFERENCE') or getattr(report, 'document_reference', None)),
        ('Revision', context.get('REVISION') or getattr(report, 'version', None)),
        ('Prepared By', context.get('PREPARED_BY') or report.prepared_by),
        ('Prepared By Email', context.get('USER_EMAIL') or report.user_email),
    ]

    prepared_on = context.get('DATE')
    if not prepared_on and isinstance(getattr(report, 'created_at', None), datetime):
        prepared_on = report.created_at.strftime('%Y-%m-%d')
    if prepared_on:
        meta_rows.append(('Prepared On', prepared_on))

    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.style = 'Table Grid'
    for label, value in meta_rows:
        if not value:
            continue
        row_cells = meta_table.add_row().cells
        _write_cell(row_cells[0], label, bold=True)
        _write_cell(row_cells[1], value)

    doc.add_paragraph('')


def _write_cell(cell, value: Any, *, bold: bool = False) -> None:
    cell.text = ''
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(_clean_text(value))
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.bold = bold


def _add_summary_sections(doc: Document, context: Dict[str, Any]) -> None:
    doc.add_heading('Summary', level=1)
    _add_text_section(doc, 'Purpose', context.get('PURPOSE'))
    _add_text_section(doc, 'Scope', context.get('SCOPE'))
    if context.get('REVISION_DETAILS'):
        _add_text_section(doc, 'Revision Details', context.get('REVISION_DETAILS'))


def _add_text_section(doc: Document, heading: str, value: Any) -> None:
    doc.add_heading(heading, level=2)
    paragraph = doc.add_paragraph(_clean_text(value) if value else 'Not provided.')
    paragraph.style = 'Normal'


def _add_table_sections(doc: Document, context: Dict[str, Any]) -> None:
    for key, title in TABLE_SECTION_KEYS:
        headers, rows = _prepare_rows(context.get(key))
        if not rows:
            continue
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.autofit = True

        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            _write_cell(header_cells[idx], header, bold=True)

        for row in rows:
            row_cells = table.add_row().cells
            for idx, header in enumerate(headers):
                _write_cell(row_cells[idx], row.get(header, ''))

        doc.add_paragraph('')


def _prepare_rows(data: Any) -> Tuple[List[str], List[Dict[str, str]]]:
    if not isinstance(data, list):
        return [], []

    headers: List[str] = []
    prepared: List[Dict[str, str]] = []

    for entry in data:
        if not isinstance(entry, dict):
            continue
        cleaned: Dict[str, str] = {}
        has_value = False
        for key, value in entry.items():
            header = _prettify_header(key)
            text = _clean_text(value)
            cleaned[header] = text
            if header not in headers:
                headers.append(header)
            if text:
                has_value = True
        if has_value:
            prepared.append(cleaned)

    return headers, prepared


def _prettify_header(value: Any) -> str:
    if not isinstance(value, str):
        return _clean_text(value)
    cleaned = value.strip()
    if '_' in cleaned:
        cleaned = cleaned.replace('_', ' ')
    cleaned = re.sub(r'\s+', ' ', cleaned)
    if cleaned.isupper() and len(cleaned) <= 4:
        return cleaned
    return cleaned.title()


def _clean_text(value: Any) -> str:
    if value is None:
        return ''
    if isinstance(value, float):
        return str(int(value)) if value.is_integer() else str(value)
    return str(value).strip()


def _build_download_name(report: Report, context: Dict[str, Any], submission_id: str) -> str:
    candidate = context.get('PROJECT_REFERENCE') or report.project_reference or context.get('DOCUMENT_TITLE') or submission_id
    safe = re.sub(r'[^A-Za-z0-9_-]+', '_', str(candidate))
    safe = safe.strip('_') or submission_id
    return f'SAT_{safe}_Modern.docx'
