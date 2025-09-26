from docx import Document
from flask import current_app

def build_sat_report(context, output_path):
    """Builds a complete .docx report programmatically based on the SAT.html form structure."""
    try:
        doc = Document('templates/SAT_Template.docx')

        # Clear placeholder content from the template
        for p in doc.paragraphs:
            p._p.getparent().remove(p._p)
        for t in doc.tables:
            t._tbl.getparent().remove(t._tbl)

        # --- 1. Document Title ---
        doc.add_heading(context.get('document_title', 'SAT Report'), level=0).style = 'Title'

        # --- 2. Static Tables ---
        def add_static_table(heading, fields):
            doc.add_heading(heading, level=1).style = 'Heading 1'
            table = doc.add_table(rows=len(fields), cols=2, style='Table Grid')
            for i, field_key in enumerate(fields):
                table.cell(i, 0).text = field_key.replace('_', ' ').title()
                table.cell(i, 1).text = str(context.get(field_key.lower(), '')) # Use lowercase keys from form

        add_static_table('Document Information', ['PROJECT_REFERENCE', 'DOCUMENT_REFERENCE', 'DATE', 'CLIENT_NAME', 'REVISION'])
        add_static_table('Approval Signatures', ['PREPARED_BY', 'REVIEWED_BY_TECH_LEAD', 'REVIEWED_BY_PM', 'APPROVED_BY_CLIENT'])
        add_static_table('Revision History', ['REVISION', 'REVISION_DETAILS', 'REVISION_DATE'])

        # --- 3. Purpose and Scope (Rich Text) ---
        doc.add_heading('Purpose', level=1).style = 'Heading 1'
        doc.add_paragraph(str(context.get('purpose', '')), style='Normal')
        doc.add_heading('Scope', level=1).style = 'Heading 1'
        doc.add_paragraph(str(context.get('scope', '')), style='Normal')

        # --- 4. Dynamic (Repeating) Tables ---
        def create_repeating_table(heading, headers, data_key_prefix):
            # Data from the form comes in as lists, e.g., doc_ref[], doc_title[]
            # We find the first list to determine the number of rows.
            first_key = f"{data_key_prefix[0]}[]"
            if first_key in context:
                num_rows = len(context[first_key])
                if num_rows > 0:
                    doc.add_heading(heading, level=1).style = 'Heading 1'
                    table = doc.add_table(rows=1, cols=len(headers), style='Table Grid')
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        hdr_cells[i].text = header
                    
                    for i in range(num_rows):
                        row_cells = table.add_row().cells
                        for j, header in enumerate(headers):
                            list_key = f"{data_key_prefix[j]}[]"
                            if list_key in context and i < len(context[list_key]):
                                row_cells[j].text = str(context[list_key][i])

        # Map headers to the 'name' attribute from the HTML form inputs
        create_repeating_table('Related Documents', ['Document Reference', 'Document Title'], ['doc_ref', 'doc_title'])
        create_repeating_table('SAT Protocol Pre-Execution Approval', ['Print Name', 'Signature', 'Date', 'Initial', 'Company'], ['pre_approval_print_name', 'pre_approval_signature', 'pre_approval_date', 'pre_approval_initial', 'pre_approval_company'])
        create_repeating_table('SAT Protocol Post Execution Approval', ['Print Name', 'Signature', 'Date', 'Initial', 'Company'], ['post_approval_print_name', 'post_approval_signature', 'post_approval_date', 'post_approval_initial', 'post_approval_company'])
        create_repeating_table('Pre-Test Requirements', ['Item', 'Test', 'Method/Test', 'Acceptance', 'Result', 'Punch', 'Verified', 'Comment'], ['pretest_item', 'pretest_test', 'pretest_method', 'pretest_acceptance', 'pretest_result', 'pretest_punch', 'pretest_verified_by', 'pretest_comment'])
        create_repeating_table('Key Components', ['S. No.', 'Model', 'Description', 'Remarks'], ['component_sno', 'component_model', 'component_description', 'component_remarks'])
        create_repeating_table('IP Address Records', ['Device Name', 'IP Address', 'Comment'], ['ip_device', 'ip_address', 'ip_comment'])
        create_repeating_table('Digital Signals', ['S.No', 'Rack', 'Pos', 'Signal TAG', 'Description', 'Result', 'Punch', 'Verified', 'Comment'], ['digital_s_no', 'digital_rack', 'digital_pos', 'digital_signal_tag', 'digital_description', 'digital_result', 'digital_punch', 'digital_verified', 'digital_comment'])
        create_repeating_table('Analogue Input Signals', ['S.No', 'Rack No', 'Module Position', 'Signal TAG', 'Description', 'Result', 'Punch Item', 'Verified By', 'Comment'], ['analogue_input_s_no', 'analogue_input_rack_no', 'analogue_input_module_position', 'analogue_input_signal_tag', 'analogue_input_description', 'analogue_input_result', 'analogue_input_punch_item', 'analogue_input_verified_by', 'analogue_input_comment'])
        create_repeating_table('Analogue Output Signals', ['S.No', 'Rack No', 'Module Position', 'Signal TAG', 'Description', 'Result', 'Punch Item', 'Verified By', 'Comment'], ['analogue_output_s_no', 'analogue_output_rack_no', 'analogue_output_module_position', 'analogue_output_signal_tag', 'analogue_output_description', 'analogue_output_result', 'analogue_output_punch_item', 'analogue_output_verified_by', 'analogue_output_comment'])
        create_repeating_table('Digital Output Signals', ['S.No', 'Rack No', 'Module Position', 'Signal TAG', 'Description', 'Result', 'Punch Item', 'Verified By', 'Comment'], ['digital_output_s_no', 'digital_output_rack_no', 'digital_output_module_position', 'digital_output_signal_tag', 'digital_output_description', 'digital_output_result', 'digital_output_punch_item', 'digital_output_verified_by', 'digital_output_comment'])
        create_repeating_table('Modbus Digital Signals', ['Address', 'Description', 'Remarks', 'Result', 'Punch Item', 'Verified By', 'Comment'], ['modbus_digital_address', 'modbus_digital_description', 'modbus_digital_remarks', 'modbus_digital_result', 'modbus_digital_punch_item', 'modbus_digital_verified_by', 'modbus_digital_comment'])
        create_repeating_table('Modbus Analogue Signals', ['Address', 'Description', 'Range', 'Result', 'Punch Item', 'Verified By', 'Comment'], ['modbus_analogue_address', 'modbus_analogue_description', 'modbus_analogue_range', 'modbus_analogue_result', 'modbus_analogue_punch_item', 'modbus_analogue_verified_by', 'modbus_analogue_comment'])
        create_repeating_table('Process Tests', ['Item', 'Action', 'Expected Result', 'Pass/Fail', 'Comments'], ['Process_Item', 'Process_Action', 'Process_Expected / Required Result', 'Process_Pass/Fail', 'Process_Comments'])
        create_repeating_table('SCADA Verification', ['Task', 'Expected Result', 'Pass/Fail', 'Comments'], ['SCADA_Task', 'SCADA_Expected_Result', 'SCADA_Pass/Fail', 'SCADA_Comments'])
        create_repeating_table('Trends Testing', ['Trend', 'Expected Behavior', 'Pass/Fail', 'Comments'], ['Trend', 'Expected Behavior', 'Pass/Fail Trend', 'Comments Trend'])

        # --- 5. Save Document ---
        doc.save(output_path)
        current_app.logger.info(f"Report built successfully at {output_path}")
        return True

    except Exception as e:
        current_app.logger.error(f"Error building report: {e}", exc_info=True)
        return False