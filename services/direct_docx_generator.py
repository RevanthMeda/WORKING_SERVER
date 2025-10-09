"""
Direct DOCX Generator - Creates SAT reports directly from database data
without using complex templates. Generates clean, properly formatted DOCX files.
"""

import os
import json
from datetime import datetime
from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from flask import current_app


class DirectSATDocxGenerator:
    """Generates SAT reports directly as DOCX files"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_document_styles()
    
    def _setup_document_styles(self):
        """Setup document styles and formatting"""
        # Set document margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        self._enable_field_updates()

    def _enable_field_updates(self):
        """Ensure Word refreshes fields (including ToC) when the document opens."""
        settings_elem = self.doc.settings.element
        update_fields = settings_elem.find(qn('w:updateFields'))
        if update_fields is None:
            update_fields = OxmlElement('w:updateFields')
            settings_elem.append(update_fields)
        update_fields.set(qn('w:val'), 'true')

    @staticmethod
    def _append_field(paragraph, instruction: str):
        """Append a Word field with the given instruction text to a paragraph."""
        # Begin field
        run = paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_char_begin)

        # Instruction text
        instr_run = paragraph.add_run()
        instr_text = OxmlElement('w:instrText')
        instr_text.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        instr_text.text = instruction
        instr_run._r.append(instr_text)

        # Separate
        run_sep = paragraph.add_run()
        fld_char_sep = OxmlElement('w:fldChar')
        fld_char_sep.set(qn('w:fldCharType'), 'separate')
        run_sep._r.append(fld_char_sep)

        # Placeholder result run (Word will replace on update)
        paragraph.add_run()

        # End field
        run_end = paragraph.add_run()
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run_end._r.append(fld_char_end)
    
    def generate_sat_report(self, report_data: Dict[str, Any], output_path: str) -> bool:
        """
        Generate a complete SAT report DOCX file
        
        Args:
            report_data: Dictionary containing all report data
            output_path: Path where to save the generated DOCX file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            current_app.logger.info("Starting direct DOCX generation")
            
            # Extract context data
            context = report_data.get('context', report_data)
            
            # Generate document sections
            self._add_header_section(context)
            self._add_document_info(context)
            self._add_approvals_section(context)
            self._add_version_control(context)
            self._add_confidentiality_notice()
            self._add_table_of_contents()
            self._add_introduction_section(context)
            self._add_pre_test_requirements(context)
            self._add_asset_register(context)
            self._add_signal_tests(context)
            self._add_process_test(context)
            self._add_scada_verification(context)
            self._add_trends_testing(context)
            self._add_alarms_section(context)
            self._add_test_equipment()
            self._add_punch_list()
            
            # Save the document
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            self.doc.save(output_path)
            
            current_app.logger.info(f"SAT report generated successfully: {output_path}")
            return True
            
        except Exception as e:
            current_app.logger.error(f"Error generating SAT report: {e}", exc_info=True)
            return False
    
    def _add_header_section(self, context: Dict[str, Any]):
        """Add document header with title and basic info"""
        # Main title
        title = self.doc.add_heading('Site Acceptance Test Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some spacing
        self.doc.add_paragraph()
    
    def _add_document_info(self, context: Dict[str, Any]):
        """Add document information section"""
        self.doc.add_heading('Document Information', level=1)
        
        # Create table for document info
        table = self.doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        # Fill document information
        info_data = [
            ('Document Title', context.get('DOCUMENT_TITLE', '')),
            ('Project Reference', context.get('PROJECT_REFERENCE', '')),
            ('Document Reference', context.get('DOCUMENT_REFERENCE', '')),
            ('Date', context.get('DATE', datetime.now().strftime('%Y-%m-%d'))),
            ('Prepared for', context.get('CLIENT_NAME', '')),
            ('Revision', context.get('REVISION', '1.0'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
        
        self.doc.add_paragraph()
    
    def _add_approvals_section(self, context: Dict[str, Any]):
        """Add document approvals section"""
        self.doc.add_heading('Document Approvals', level=1)
        
        # Create approvals table
        table = self.doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'
        
        # Headers
        headers = ['Role', 'Name', 'Date']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        
        # Approval data
        approvals_data = [
            ('Prepared by', context.get('PREPARED_BY', ''), context.get('PREPARER_DATE', '')),
            ('Reviewed by (Tech Lead)', context.get('REVIEWED_BY_TECH_LEAD', ''), context.get('TECH_LEAD_DATE', '')),
            ('Reviewed by (PM)', context.get('REVIEWED_BY_PM', ''), context.get('PM_DATE', '')),
            ('Approved by (Client)', context.get('APPROVED_BY_CLIENT', ''), '')
        ]
        
        for i, (role, name, date) in enumerate(approvals_data, 1):
            table.cell(i, 0).text = role
            table.cell(i, 1).text = str(name)
            table.cell(i, 2).text = str(date)
        
        self.doc.add_paragraph()
    
    def _add_version_control(self, context: Dict[str, Any]):
        """Add version control section"""
        self.doc.add_heading('Document Version Control', level=1)
        
        table = self.doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        
        # Headers
        headers = ['Revision Number', 'Details', 'Date']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        
        # Version data
        table.cell(1, 0).text = context.get('REVISION', '1.0')
        table.cell(1, 1).text = context.get('REVISION_DETAILS', 'Initial version')
        table.cell(1, 2).text = context.get('REVISION_DATE', datetime.now().strftime('%Y-%m-%d'))
        
        self.doc.add_paragraph()
    
    def _add_confidentiality_notice(self):
        """Add confidentiality notice"""
        para = self.doc.add_paragraph()
        para.add_run('Confidentiality Notice: ').bold = True
        para.add_run('This document contains confidential and proprietary information of Cully. '
                    'Unauthorized distribution or reproduction is strictly prohibited.')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add Cully website
        website_para = self.doc.add_paragraph('WWW.CULLY.IE')
        website_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        website_para.runs[0].bold = True
        
        self.doc.add_page_break()
    
    def _add_table_of_contents(self):
        """Add table of contents"""
        self.doc.add_heading('Table of Contents', level=1)
        toc_paragraph = self.doc.add_paragraph()
        self._append_field(toc_paragraph, r'TOC \o "1-3" \h \z \u')

        note = self.doc.add_paragraph()
        note_style_name = 'Intense Quote'
        if note_style_name in [style.name for style in self.doc.styles]:
            note.style = note_style_name
        note_run = note.add_run('Hint: Right-click the table of contents and select "Update Field" after opening the document to refresh page numbers.')
        note_run.italic = True

        total_para = self.doc.add_paragraph()
        total_label = total_para.add_run('Total pages: ')
        total_label.bold = True
        self._append_field(total_para, 'NUMPAGES')

        self.doc.add_page_break()
    
    def _add_introduction_section(self, context: Dict[str, Any]):
        """Add introduction section"""
        self.doc.add_heading('1. Introduction', level=1)
        
        # Purpose
        self.doc.add_heading('1.1 Purpose', level=2)
        self.doc.add_paragraph(context.get('PURPOSE', ''))
        
        # Scope
        self.doc.add_heading('1.2 Scope', level=2)
        self.doc.add_paragraph(context.get('SCOPE', ''))
        
        # Related documents
        self.doc.add_heading('1.3 Relationship with Other Documents', level=2)
        related_docs = context.get('RELATED_DOCUMENTS', [])
        if related_docs:
            table = self.doc.add_table(rows=len(related_docs) + 1, cols=2)
            table.style = 'Table Grid'
            table.cell(0, 0).text = 'Document Reference'
            table.cell(0, 1).text = 'Document Title'
            
            for i, doc in enumerate(related_docs, 1):
                table.cell(i, 0).text = doc.get('Document_Reference', '')
                table.cell(i, 1).text = doc.get('Document_Title', '')
        
        self.doc.add_paragraph()
    
    def _add_pre_test_requirements(self, context: Dict[str, Any]):
        """Add pre-test requirements section"""
        self.doc.add_heading('2. Pre-Test Requirements', level=1)
        
        pre_test_data = context.get('PRE_TEST_REQUIREMENTS', [])
        if pre_test_data:
            table = self.doc.add_table(rows=len(pre_test_data) + 1, cols=8)
            table.style = 'Table Grid'
            
            # Headers
            headers = ['Item', 'Test', 'Method/Test Steps', 'Acceptance Criteria', 
                      'Result', 'Punch Item', 'Verified by', 'Comment']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
            
            # Data rows
            for i, row_data in enumerate(pre_test_data, 1):
                for j, header in enumerate(headers):
                    table.cell(i, j).text = str(row_data.get(header, ''))
        
        self.doc.add_paragraph()
    
    def _add_asset_register(self, context: Dict[str, Any]):
        """Add asset register section"""
        self.doc.add_heading('3. Asset Register', level=1)
        
        # Key Components
        self.doc.add_heading('3.1 Key Components', level=2)
        key_components = context.get('KEY_COMPONENTS', [])
        if key_components:
            table = self.doc.add_table(rows=len(key_components) + 1, cols=4)
            table.style = 'Table Grid'
            
            headers = ['S.no', 'Model', 'Description', 'Remarks']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
            
            for i, component in enumerate(key_components, 1):
                table.cell(i, 0).text = str(component.get('S_no', ''))
                table.cell(i, 1).text = str(component.get('Model', ''))
                table.cell(i, 2).text = str(component.get('Description', ''))
                table.cell(i, 3).text = str(component.get('Remarks', ''))
        
        # IP Address Records
        self.doc.add_heading('3.2 IP Address Records', level=2)
        ip_records = context.get('IP_RECORDS', [])
        if ip_records:
            table = self.doc.add_table(rows=len(ip_records) + 1, cols=3)
            table.style = 'Table Grid'
            
            headers = ['Device', 'IP Address/Subnet/Gateway', 'Comments']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
            
            for i, record in enumerate(ip_records, 1):
                table.cell(i, 0).text = str(record.get('Device_Name', ''))
                table.cell(i, 1).text = str(record.get('IP_Address', ''))
                table.cell(i, 2).text = str(record.get('Comment', ''))
        
        self.doc.add_paragraph()
    
    def _add_signal_tests(self, context: Dict[str, Any]):
        """Add signal tests section"""
        self.doc.add_heading('4. Signal Tests', level=1)
        
        # Digital Input Signals
        self._add_signal_table('4.1 Digital Input Signals', context.get('SIGNAL_LISTS', []))
        
        # Digital Output Signals
        self._add_signal_table('4.2 Digital Output Signals', context.get('DIGITAL_OUTPUT_LISTS', []))
        
        # Analogue Input Signals
        self._add_signal_table('4.3 Analogue Input Signals', context.get('ANALOGUE_LISTS', []))
        
        # Analogue Output Signals
        self._add_signal_table('4.4 Analogue Output Signals', context.get('ANALOGUE_OUTPUT_LISTS', []))
        
        # Modbus Digital
        self._add_modbus_digital_table(context.get('MODBUS_DIGITAL_LISTS', []))
        
        # Modbus Analogue
        self._add_modbus_analogue_table(context.get('MODBUS_ANALOGUE_LISTS', []))
        
        # Data Validation
        self._add_data_validation_table(context.get('DATA_VALIDATION', []))
    
    def _add_signal_table(self, title: str, signal_data: List[Dict]):
        """Add a signal test table"""
        self.doc.add_heading(title, level=2)
        
        if signal_data:
            table = self.doc.add_table(rows=len(signal_data) + 1, cols=9)
            table.style = 'Table Grid'
            
            headers = ['S. No.', 'Rack No.', 'Module Position', 'Signal TAG', 
                      'Signal Description', 'Result', 'Punch Item', 'Verified By', 'Comment']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
            
            for i, signal in enumerate(signal_data, 1):
                for j, header in enumerate(headers):
                    key = header.replace('.', '')  # Remove dots for key matching
                    table.cell(i, j).text = str(signal.get(key, ''))
        
        self.doc.add_paragraph()
    
    def _add_modbus_digital_table(self, modbus_data: List[Dict]):
        """Add Modbus digital table"""
        self.doc.add_heading('4.5 Modbus Digital', level=2)
        
        if modbus_data:
            table = self.doc.add_table(rows=len(modbus_data) + 1, cols=7)
            table.style = 'Table Grid'
            
            headers = ['Address', 'Description', 'Remarks', 'Result', 'Punch Item', 'Verified By', 'Comment']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
            
            for i, data in enumerate(modbus_data, 1):
                for j, header in enumerate(headers):
                    table.cell(i, j).text = str(data.get(header, ''))
        
        self.doc.add_paragraph()
    
    def _add_modbus_analogue_table(self, modbus_data: List[Dict]):
        """Add Modbus analogue table"""
        self.doc.add_heading('4.6 Modbus Analogue', level=2)
        
        if modbus_data:
            table = self.doc.add_table(rows=len(modbus_data) + 1, cols=7)
            table.style = 'Table Grid'
            
            headers = ['Address', 'Description', 'Range', 'Result', 'Punch Item', 'Verified By', 'Comment']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
            
            for i, data in enumerate(modbus_data, 1):
                for j, header in enumerate(headers):
                    # Handle the space in ' Address' key
                    key = ' Address' if header == 'Address' else header
                    table.cell(i, j).text = str(data.get(key, ''))
        
        self.doc.add_paragraph()
    
    def _add_data_validation_table(self, validation_data: List[Dict]):
        """Add data validation table"""
        self.doc.add_heading('4.7 Data Validation', level=2)
        
        if validation_data:
            table = self.doc.add_table(rows=len(validation_data) + 1, cols=4)
            table.style = 'Table Grid'
            
            headers = ['Tag', 'Range', 'SCADA Value', 'HMI Value']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
            
            for i, data in enumerate(validation_data, 1):
                for j, header in enumerate(headers):
                    table.cell(i, j).text = str(data.get(header, ''))
        
        self.doc.add_paragraph()
    
    def _add_process_test(self, context: Dict[str, Any]):
        """Add process test section"""
        self.doc.add_heading('5. Process Test', level=1)
        
        process_data = context.get('PROCESS_TEST', [])
        if process_data:
            table = self.doc.add_table(rows=len(process_data) + 1, cols=5)
            table.style = 'Table Grid'
            
            headers = ['Item', 'Action', 'Expected / Required Result', 'Pass/Fail', 'Comments']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
            
            for i, data in enumerate(process_data, 1):
                for j, header in enumerate(headers):
                    # Handle spaces in keys
                    key = ' Pass/Fail ' if 'Pass/Fail' in header else (' Comments ' if 'Comments' in header else header)
                    table.cell(i, j).text = str(data.get(key, ''))
        
        self.doc.add_paragraph()
    
    def _add_scada_verification(self, context: Dict[str, Any]):
        """Add SCADA verification section"""
        self.doc.add_heading('6. SCADA Verification', level=1)
        
        scada_data = context.get('SCADA_VERIFICATION', [])
        if scada_data:
            table = self.doc.add_table(rows=len(scada_data) + 1, cols=4)
            table.style = 'Table Grid'
            
            headers = ['Task', 'Expected Result', 'Pass/Fail', 'Comments']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
            
            for i, data in enumerate(scada_data, 1):
                for j, header in enumerate(headers):
                    table.cell(i, j).text = str(data.get(header, ''))
        
        # Add note about images
        self.doc.add_paragraph('SCADA screenshots and verification images would be inserted here.')
        self.doc.add_paragraph()
    
    def _add_trends_testing(self, context: Dict[str, Any]):
        """Add trends testing section"""
        self.doc.add_heading('7. Trends Testing', level=1)
        
        trends_data = context.get('TRENDS_TESTING', [])
        if trends_data:
            table = self.doc.add_table(rows=len(trends_data) + 1, cols=4)
            table.style = 'Table Grid'
            
            headers = ['Trend', 'Expected Behavior', 'Pass/Fail', 'Comments']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
            
            for i, data in enumerate(trends_data, 1):
                for j, header in enumerate(headers):
                    table.cell(i, j).text = str(data.get(header, ''))
        
        # Add note about images
        self.doc.add_paragraph('Trends screenshots and analysis charts would be inserted here.')
        self.doc.add_paragraph()
    
    def _add_alarms_section(self, context: Dict[str, Any]):
        """Add alarms section"""
        self.doc.add_heading('8. SCADA / SMS Alarms', level=1)
        
        alarm_data = context.get('ALARM_LIST', [])
        if alarm_data:
            table = self.doc.add_table(rows=len(alarm_data) + 1, cols=4)
            table.style = 'Table Grid'
            
            headers = ['Alarm Type', 'Expected / Required Result', 'Pass/Fail', 'Comments']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
            
            for i, data in enumerate(alarm_data, 1):
                for j, header in enumerate(headers):
                    # Handle spaces in keys
                    key = ' Pass/Fail ' if 'Pass/Fail' in header else (' Comments ' if 'Comments' in header else header)
                    table.cell(i, j).text = str(data.get(key, ''))
        
        # Add note about images
        self.doc.add_paragraph('Alarm screenshots and notification examples would be inserted here.')
        self.doc.add_paragraph()
    
    def _add_test_equipment(self):
        """Add test equipment section"""
        self.doc.add_heading('9. List of Test Equipment', level=1)
        
        table = self.doc.add_table(rows=8, cols=5)
        table.style = 'Table Grid'
        
        headers = ['Item', 'Equipment/Tool', 'Make/Model', 'Serial No.', 'Calibration Date']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        
        # Add numbered rows
        for i in range(1, 8):
            table.cell(i, 0).text = str(i)
        
        self.doc.add_paragraph()
    
    def _add_punch_list(self):
        """Add punch list section"""
        self.doc.add_heading('10. Punch List Items', level=1)
        
        table = self.doc.add_table(rows=12, cols=6)
        table.style = 'Table Grid'
        
        headers = ['Item', 'Detail of Non-Conformance', 'Raised By', 'Comment', 'Approved By', 'Date']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        
        self.doc.add_paragraph()


def generate_sat_report_direct(submission_id: str) -> Dict[str, Any]:
    """
    Generate SAT report directly from database data
    
    Args:
        submission_id: The report ID
        
    Returns:
        Dict with 'path' and 'download_name' on success, or 'error' on failure
    """
    try:
        from models import Report, SATReport
        
        # Load report from database
        report = Report.query.filter_by(id=submission_id).first()
        if not report:
            return {'error': 'Report not found'}

        sat_report = SATReport.query.filter_by(report_id=submission_id).first()
        if not sat_report:
            return {'error': 'Report data not found'}

        # Parse stored data
        try:
            stored_data = json.loads(sat_report.data_json) if sat_report.data_json else {}
        except json.JSONDecodeError:
            current_app.logger.error(f"Invalid JSON in report {submission_id}")
            return {'error': 'Invalid report data'}

        context_data = stored_data.get('context', stored_data) or {}
        
        # Add basic report info if missing
        if not context_data.get('DOCUMENT_TITLE'):
            context_data['DOCUMENT_TITLE'] = report.document_title or 'SAT Report'
        if not context_data.get('CLIENT_NAME'):
            context_data['CLIENT_NAME'] = report.client_name or ''
        if not context_data.get('PROJECT_REFERENCE'):
            context_data['PROJECT_REFERENCE'] = report.project_reference or ''
        
        # Generate output path
        output_dir = current_app.config.get('OUTPUT_DIR', 'outputs')
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, f'SAT_Direct_{submission_id}.docx')
        
        # Generate the document
        generator = DirectSATDocxGenerator()
        success = generator.generate_sat_report(context_data, output_path)
        
        if success:
            # Build download name
            project_ref = context_data.get('PROJECT_REFERENCE', submission_id[:8])
            safe_proj_ref = "".join(c if c.isalnum() or c in ['_', '-'] else "_" for c in project_ref)
            download_name = f"SAT_{safe_proj_ref}.docx"
            
            return {
                'path': output_path,
                'download_name': download_name
            }
        else:
            return {'error': 'Failed to generate document'}
            
    except Exception as e:
        current_app.logger.error(f"Error in generate_sat_report_direct: {e}", exc_info=True)
        return {'error': str(e)}
