"""
Service for regenerating SAT reports from database data with full template support
"""
import os
import json
import tempfile
import datetime as dt
import re
from html import unescape
from typing import Dict, Any, List
from flask import current_app, url_for
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

from models import Report, SATReport, User
from services.sat_tables import extract_ui_tables, build_doc_tables, migrate_context_tables, TABLE_CONFIG
from utils import update_toc_page_numbers


HTML_TAG_RE = re.compile(r'<[^>]+>')
BR_RE = re.compile(r'<br\s*/?>', re.IGNORECASE)
PARA_CLOSE_RE = re.compile(r'</p\s*>', re.IGNORECASE)
PARA_OPEN_RE = re.compile(r'<p\s*>', re.IGNORECASE)
LI_OPEN_RE = re.compile(r'<li\s*>', re.IGNORECASE)
LI_CLOSE_RE = re.compile(r'</li\s*>', re.IGNORECASE)


def _strip_html(value: str) -> str:
    """Convert simple HTML snippets to plain text for Word templates."""
    text = value.replace('\r\n', '\n')
    text = BR_RE.sub('\n', text)
    text = PARA_CLOSE_RE.sub('\n', text)
    text = PARA_OPEN_RE.sub('', text)
    text = LI_CLOSE_RE.sub('\n', text)
    text = LI_OPEN_RE.sub('- ', text)
    text = HTML_TAG_RE.sub('', text)
    text = unescape(text)
    text = text.replace('\xa0', ' ')
    lines = [line.rstrip() for line in text.splitlines()]
    cleaned = '\n'.join(line for line in lines if line)
    return cleaned.strip()


SIGNATURE_WIDTH_MM = 40


def _format_timestamp(value: str) -> str:
    """Convert stored timestamps into a readable string for templates."""
    if not value:
        return ""
    cleaned = value.strip() if isinstance(value, str) else str(value)
    if not cleaned:
        return ""
    try:
        return dt.datetime.fromisoformat(cleaned).strftime('%Y-%m-%d %H:%M')
    except Exception:
        return cleaned


def _load_signature_image(doc: DocxTemplate, value: Any) -> Any:
    """Attempt to rebuild an InlineImage for stored signature references."""
    if not value:
        return ""
    if isinstance(value, InlineImage):
        return value

    filename = str(value).strip()
    if not filename:
        return ""

    candidates = []
    base_names = [filename]
    if not os.path.splitext(filename)[1]:
        base_names.append(f"{filename}.png")

    # Always try the stored value directly (covers relative paths with separators)
    for name in base_names:
        candidates.append(name)

    if os.path.isabs(filename):
        candidates.append(filename)
    else:
        base_dirs = [
            current_app.config.get('SIGNATURES_FOLDER'),
            os.path.join(current_app.root_path, 'static', 'signatures'),
            os.path.join(os.getcwd(), 'static', 'signatures'),
        ]

        for base in filter(None, base_dirs):
            for name in base_names:
                candidates.append(os.path.join(base, name))

    for candidate in candidates:
        try:
            if os.path.exists(candidate) and os.path.getsize(candidate) > 0:
                return InlineImage(doc, candidate, width=Mm(SIGNATURE_WIDTH_MM))
        except Exception as exc:
            current_app.logger.error(
                f"Failed to load signature image from {candidate}: {exc}",
                exc_info=True
            )
    return ""


def _select_stage_approval(approvals: List[Dict[str, Any]], stage: int) -> Dict[str, Any]:
    """Return the approval entry that matches the provided stage number."""
    for approval in approvals or []:
        try:
            if int(approval.get('stage', 0)) == int(stage):
                return approval
        except (TypeError, ValueError):
            continue
    return {}


def _append_word_field(paragraph, instruction: str) -> None:
    """Append a Word field code to the given paragraph."""
    run_begin = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    run_begin._r.append(fld_char_begin)

    instr_run = paragraph.add_run()
    instr_text = OxmlElement('w:instrText')
    instr_text.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr_text.text = instruction
    instr_run._r.append(instr_text)

    run_sep = paragraph.add_run()
    fld_char_sep = OxmlElement('w:fldChar')
    fld_char_sep.set(qn('w:fldCharType'), 'separate')
    run_sep._r.append(fld_char_sep)

    paragraph.add_run()  # placeholder result

    run_end = paragraph.add_run()
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run_end._r.append(fld_char_end)


def _enable_field_updates(document: DocxTemplate) -> None:
    """Ensure Word refreshes fields (e.g., TOC) when the document is opened."""
    try:
        settings_elem = document.settings.element
        update_fields = settings_elem.find(qn('w:updateFields'))
        if update_fields is None:
            update_fields = OxmlElement('w:updateFields')
            settings_elem.append(update_fields)
        update_fields.set(qn('w:val'), 'true')
    except Exception as exc:
        current_app.logger.warning(f"Could not enable field updates in document: {exc}")


def _ensure_total_pages_line(document: DocxTemplate) -> None:
    """Insert a total pages line near the table of contents if not present."""
    target_paragraph = None
    for paragraph in document.paragraphs:
        if paragraph.text.strip().lower() == 'table of contents':
            target_paragraph = paragraph
            break

    if target_paragraph:
        insert_after = target_paragraph.insert_paragraph_after('Total pages: ')
    else:
        insert_after = document.add_paragraph('Total pages: ')

    if insert_after.runs:
        insert_after.runs[0].bold = True
    _append_word_field(insert_after, 'NUMPAGES')


def regenerate_document_from_db(submission_id: str) -> Dict[str, Any]:
    """
    Regenerate a SAT report document from database data
    Returns dict with 'path' and 'download_name' on success, or 'error' on failure
    """
    try:
        # Load report from database
        report = Report.query.filter_by(id=submission_id).first()
        if not report:
            return {'error': 'Report not found'}

        sat_report = SATReport.query.filter_by(report_id=submission_id).first()
        if not sat_report:
            return {'error': 'Report data not found'}

        # Parse stored data
        try:
            stored_data = json.loads(sat_report.data_json) if sat_report.data_json else {}
        except json.JSONDecodeError:
            current_app.logger.error(f"Invalid JSON in report {submission_id}")
            return {'error': 'Invalid report data'}

        context_data = stored_data.get('context', stored_data) or {}

        approvals_data: List[Dict[str, Any]] = []
        if report.approvals_json:
            try:
                approvals_data = json.loads(report.approvals_json) or []
            except json.JSONDecodeError:
                current_app.logger.warning(
                    "Unable to decode approvals JSON for report %s",
                    submission_id
                )
                approvals_data = []

        tech_approval = _select_stage_approval(approvals_data, 1)
        pm_approval = _select_stage_approval(approvals_data, 2)
        client_approval = _select_stage_approval(approvals_data, 3)

        # Load template
        template_path = current_app.config.get('TEMPLATE_FILE')
        if not template_path or not os.path.exists(template_path):
            return {'error': 'Template file not found'}

        doc = DocxTemplate(template_path)
        current_app.logger.info(f"Template loaded for regeneration: {submission_id}")

        # Load existing images and convert to InlineImage objects
        scada_urls = json.loads(sat_report.scada_image_urls) if sat_report.scada_image_urls else []
        trends_urls = json.loads(sat_report.trends_image_urls) if sat_report.trends_image_urls else []
        alarm_urls = json.loads(sat_report.alarm_image_urls) if sat_report.alarm_image_urls else []

        scada_image_objects = _load_existing_images(doc, scada_urls)
        trends_image_objects = _load_existing_images(doc, trends_urls)
        alarm_image_objects = _load_existing_images(doc, alarm_urls)

        current_app.logger.info(f"Loaded {len(scada_image_objects)} SCADA, {len(trends_image_objects)} Trends, {len(alarm_image_objects)} Alarm images")

        # Build table data - extract all UI section keys from TABLE_CONFIG
        ui_tables = {}
        for config in TABLE_CONFIG:
            table_key = config['ui_section']
            table_data = context_data.get(table_key, [])
            if table_data:
                ui_tables[table_key] = table_data

        doc_tables = build_doc_tables(ui_tables) if ui_tables else {}
        legacy_tables = migrate_context_tables(context_data)
        combined_tables = dict(legacy_tables)
        combined_tables.update(ui_tables)

        # Helper to sanitize values - convert None to empty string
        # Note: docxtpl automatically handles XML escaping, so we don't need to do it manually
        sig_prepared_source = (
            context_data.get('prepared_signature')
            or context_data.get('SIG_PREPARED')
        )
        sig_review_tech_source = (
            context_data.get('SIG_REVIEW_TECH')
            or tech_approval.get('signature')
        )
        sig_review_pm_source = (
            context_data.get('SIG_REVIEW_PM')
            or pm_approval.get('signature')
        )
        sig_client_source = (
            context_data.get('SIG_APPROVAL_CLIENT')
            or client_approval.get('signature')
        )

        sig_prepared = _load_signature_image(doc, sig_prepared_source)
        sig_review_tech = _load_signature_image(doc, sig_review_tech_source)
        sig_review_pm = _load_signature_image(doc, sig_review_pm_source)
        sig_approval_client = _load_signature_image(doc, sig_client_source)
        sig_approver_1 = _load_signature_image(doc, context_data.get('SIG_APPROVER_1'))
        sig_approver_2 = _load_signature_image(doc, context_data.get('SIG_APPROVER_2'))
        sig_approver_3 = _load_signature_image(doc, context_data.get('SIG_APPROVER_3'))

        preparer_date_value = context_data.get('PREPARER_DATE') or context_data.get('prepared_timestamp')
        tech_lead_date_value = (
            context_data.get('TECH_LEAD_DATE')
            or context_data.get('tech_lead_timestamp')
            or tech_approval.get('timestamp')
        )
        pm_date_value = (
            context_data.get('PM_DATE')
            or context_data.get('pm_timestamp')
            or pm_approval.get('timestamp')
        )

        def sanitize_value(value):
            if value is None:
                return ""
            if isinstance(value, str):
                cleaned = value.strip()
                cleaned = cleaned.replace('\xa0', ' ')
                if HTML_TAG_RE.search(cleaned):
                    cleaned = _strip_html(cleaned)
                return cleaned
            return str(value)

        def resolve_approver_name(existing_value: str, approval_entry: Dict[str, Any]) -> str:
            """Use the freshest available name for an approval stage."""
            candidate = existing_value or ""
            if approval_entry:
                candidate = approval_entry.get('approver_name') or candidate
                approver_email = approval_entry.get('approver_email')
                if approver_email:
                    user = User.query.filter_by(email=approver_email).first()
                    if user and user.full_name:
                        candidate = user.full_name
            return candidate

        tech_lead_name = resolve_approver_name(
            context_data.get('REVIEWED_BY_TECH_LEAD', ''),
            tech_approval
        )
        pm_name = resolve_approver_name(
            context_data.get('REVIEWED_BY_PM', ''),
            pm_approval
        )
        client_approver_name = resolve_approver_name(
            context_data.get('APPROVED_BY_CLIENT', ''),
            client_approval
        )

        # Build rendering context with sanitized values - include ALL template fields
        render_context = {
            "DOCUMENT_TITLE": sanitize_value(context_data.get('DOCUMENT_TITLE', '')),
            "PROJECT_REFERENCE": sanitize_value(context_data.get('PROJECT_REFERENCE', '')),
            "DOCUMENT_REFERENCE": sanitize_value(context_data.get('DOCUMENT_REFERENCE', '')),
            "DATE": sanitize_value(context_data.get('DATE', '')),
            "CLIENT_NAME": sanitize_value(context_data.get('CLIENT_NAME', '')),
            "REVISION": sanitize_value(context_data.get('REVISION', '')),
            "REVISION_DETAILS": sanitize_value(context_data.get('REVISION_DETAILS', '')),
            "REVISION_DATE": sanitize_value(context_data.get('REVISION_DATE', '')),
            "PREPARED_BY": sanitize_value(context_data.get('PREPARED_BY', '')),
            "PREPARER_DATE": sanitize_value(_format_timestamp(preparer_date_value)),
            "TECH_LEAD_DATE": sanitize_value(_format_timestamp(tech_lead_date_value)),
            "PM_DATE": sanitize_value(_format_timestamp(pm_date_value)),
            "SIG_PREPARED": sig_prepared,
            "SIG_PREPARED_BY": sanitize_value(context_data.get('SIG_PREPARED_BY', context_data.get('PREPARED_BY', ''))),
            "REVIEWED_BY_TECH_LEAD": sanitize_value(tech_lead_name),
            "SIG_REVIEW_TECH": sig_review_tech,
            "REVIEWED_BY_PM": sanitize_value(pm_name),
            "SIG_REVIEW_PM": sig_review_pm,
            "APPROVED_BY_CLIENT": sanitize_value(client_approver_name),
            "SIG_APPROVAL_CLIENT": sig_approval_client,
            "PURPOSE": sanitize_value(context_data.get('PURPOSE', '')),
            "SCOPE": sanitize_value(context_data.get('SCOPE', '')),
            "SCADA_IMAGES": scada_image_objects if scada_image_objects else [],
            "TRENDS_IMAGES": trends_image_objects if trends_image_objects else [],
            "ALARM_IMAGES": alarm_image_objects if alarm_image_objects else [],
            "SCADA_SCREENSHOTS": scada_urls if scada_urls else [],
            "TRENDS_SCREENSHOTS": trends_urls if trends_urls else [],
            "ALARM_SCREENSHOTS": alarm_urls if alarm_urls else [],
            "SIG_APPROVER_1": sig_approver_1,
            "SIG_APPROVER_2": sig_approver_2,
            "SIG_APPROVER_3": sig_approver_3,
        }

        # Add table data with sanitization
        def normalize_table_keys(data):
            """Normalize table data to ensure no None values and handle common key variations"""
            if data is None:
                return []
            if isinstance(data, list):
                normalized_list = []
                for item in data:
                    if isinstance(item, dict):
                        # Create normalized dict with sanitized values and key variations
                        normalized_dict = {}
                        for k, v in item.items():
                            value = sanitize_value(v)
                            # Add both original and trimmed versions of the key
                            normalized_dict[k] = value
                            if k != k.strip():
                                normalized_dict[k.strip()] = value
                        normalized_list.append(normalized_dict)
                    else:
                        normalized_list.append(sanitize_value(item))
                return normalized_list
            return data
        
        for key, value in doc_tables.items():
            if key not in render_context:
                render_context[key] = normalize_table_keys(value)

        # Render template
        current_app.logger.info("Starting document rendering from database...")
        doc.render(render_context)
        current_app.logger.info("Document rendering completed")
        _enable_field_updates(doc)
        _ensure_total_pages_line(doc)

        # Save to temp file
        timestamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"SAT_Report_{timestamp}.docx"
        temp_path = os.path.join(tempfile.gettempdir(), filename)
        
        doc.save(temp_path)

        # Optionally refresh TOC page numbers (Windows/Word only)
        try:
            if current_app.config.get('AUTO_UPDATE_TOC', False):
                update_toc_page_numbers(temp_path)
        except Exception as toc_error:
            current_app.logger.warning(f"TOC page-number update skipped during regeneration: {toc_error}")
        
        # Verify file
        if not os.path.exists(temp_path):
            return {'error': 'Document file was not created'}
        
        file_size = os.path.getsize(temp_path)
        if file_size < 1000:
            return {'error': f'Document file is too small ({file_size} bytes)'}
        
        current_app.logger.info(f"Document regenerated successfully: {file_size} bytes at {temp_path}")

        # Build download name
        project_ref = context_data.get('PROJECT_REFERENCE', '').strip()
        if not project_ref:
            project_ref = submission_id[:8]
        safe_proj_ref = "".join(c if c.isalnum() or c in ['_', '-'] else "_" for c in project_ref)
        download_name = f"SAT_{safe_proj_ref}.docx"

        return {
            'path': temp_path,
            'download_name': download_name
        }

    except Exception as e:
        current_app.logger.error(f"Error regenerating document for {submission_id}: {e}", exc_info=True)
        return {'error': str(e)}


def _load_existing_images(doc: DocxTemplate, url_list: List[str]) -> List[InlineImage]:
    """Convert existing image URLs to InlineImage objects for Word template"""
    image_objects = []
    for url in url_list:
        try:
            # Convert URL back to file path
            # URL format: /static/uploads/{submission_id}/{filename}
            if '/uploads/' in url:
                parts = url.split('/uploads/')
                if len(parts) == 2:
                    rel_path = parts[1]
                    disk_path = os.path.join(current_app.config['UPLOAD_ROOT'], rel_path.replace('/', os.sep))
                    
                    if os.path.exists(disk_path):
                        # Get image dimensions
                        with Image.open(disk_path) as img:
                            w, h = img.size
                        
                        # Calculate scale to fit max width
                        max_w_mm = 150
                        scale = min(1, max_w_mm / (w * 0.264583))
                        
                        # Create InlineImage
                        image_objects.append(
                            InlineImage(doc, disk_path,
                                width=Mm(w * 0.264583 * scale),
                                height=Mm(h * 0.264583 * scale)
                            )
                        )
                        current_app.logger.debug(f"Loaded existing image: {disk_path}")
                    else:
                        current_app.logger.warning(f"Image file not found: {disk_path}")
        except Exception as e:
            current_app.logger.error(f"Error loading existing image {url}: {e}", exc_info=True)
    return image_objects
