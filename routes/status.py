from flask import Blueprint, render_template, redirect, url_for, flash, current_app, Response, jsonify, request
import os
import json
from flask_login import current_user, login_required
from services.sat_tables import build_doc_tables_from_context, migrate_context_tables
import datetime as dt

status_bp = Blueprint('status', __name__)

@status_bp.route('/<submission_id>')
@login_required
def view_status(submission_id):
    """View a specific submission with auto-download"""
    from models import Report, SATReport

    # Check if submission_id is valid
    if not submission_id or submission_id == 'None':
        flash('Invalid submission ID.', 'error')
        return redirect(url_for('dashboard.home'))

    report = Report.query.filter_by(id=submission_id).first()
    if not report:
        flash('Report not found.', 'error')
        return redirect(url_for('dashboard.home'))

    sat_report = SATReport.query.filter_by(report_id=report.id).first()
    if not sat_report:
        flash('Report data not found.', 'error')
        return redirect(url_for('dashboard.home'))

    try:
        stored_data = json.loads(sat_report.data_json) if sat_report.data_json else {}
    except json.JSONDecodeError:
        stored_data = {}

    approvals = json.loads(report.approvals_json) if report.approvals_json else []
    for approval in approvals:
        flags = approval.get("flags", [])
        if not isinstance(flags, list):
            approval["flags"] = []
            continue
        cleaned_flags = []
        for item in flags:
            if not isinstance(item, dict):
                continue
            cleaned_flags.append({
                "location": (item.get("location") or "").strip(),
                "note": (item.get("note") or "").strip(),
                "action": (item.get("action") or "").strip(),
                "severity": (item.get("severity") or "Medium").strip() or "Medium",
            })
        approval["flags"] = cleaned_flags

    # Determine overall status
    statuses = [a.get("status", "pending") for a in approvals]
    if "rejected" in statuses:
        overall_status = "rejected"
    elif all(status == "approved" for status in statuses):
        overall_status = "approved"
    elif any(status == "approved" for status in statuses):
        overall_status = "partially_approved"
    else:
        overall_status = "pending"

    # Get submission data context with fallbacks
    submission_data = stored_data.get("context", {})
    if not submission_data:
        submission_data = stored_data  # Fallback if context doesn't exist

    # Check if report files exist
    pdf_path = os.path.join(current_app.config['OUTPUT_DIR'], f'SAT_Report_{submission_id}_Final.pdf')
    docx_path = os.path.join(current_app.config['OUTPUT_DIR'], f'SAT_Report_{submission_id}_Final.docx')

    download_available = os.path.exists(pdf_path) or os.path.exists(docx_path)
    has_pdf = os.path.exists(pdf_path)

    # Determine if current user can edit this report
    can_edit = False
    if current_user.role == 'Admin':
        can_edit = True  # Admin can edit any report
    elif current_user.role == 'Engineer' and current_user.email == report.user_email:
        # Engineers can edit their own reports until approved by Automation Manager
        tm_approved = any(a.get("status") == "approved" and a.get("stage") == 1 for a in approvals)
        can_edit = not tm_approved
    elif current_user.role == 'Automation Manager':
        # Automation Manager can edit reports until approved by PM
        pm_approved = any(a.get("status") == "approved" and a.get("stage") == 2 for a in approvals)
        can_edit = not pm_approved

    # Build context similar to old version
    context = {
        "submission_id": submission_id,
        "submission_data": submission_data,
        "approvals": approvals,
        "locked": report.locked,
        "can_edit": can_edit,
        "created_at": report.created_at.strftime('%Y-%m-%d %H:%M:%S') if isinstance(report.created_at, dt.datetime) else report.created_at,
        "updated_at": report.updated_at.strftime('%Y-%m-%d %H:%M:%S') if isinstance(report.updated_at, dt.datetime) else report.updated_at,
        "user_email": report.user_email,
        "document_title": submission_data.get("DOCUMENT_TITLE", "SAT Report"),
        "project_reference": submission_data.get("PROJECT_REFERENCE", ""),
        "client_name": submission_data.get("CLIENT_NAME", ""),
        "prepared_by": submission_data.get("PREPARED_BY", ""),
        "overall_status": overall_status,
        "download_available": download_available,
        "has_pdf": has_pdf,
        "auto_download": True
    }

    return render_template('status.html', **context)

@status_bp.route('/download/<submission_id>')
@login_required
def download_report(submission_id):
    """Generate and download SAT report using the official template when available"""
    try:
        current_app.logger.info(f"Starting SAT download generation for {submission_id}")
        
        from services.document_generator import regenerate_document_from_db

        # Prefer generating from the SAT template first
        result = regenerate_document_from_db(submission_id)

        if 'error' in result:
            current_app.logger.warning(
                f"Template-based generation failed for {submission_id}: {result['error']}. Falling back to direct generator."
            )
            from services.direct_docx_generator import generate_sat_report_direct
            result = generate_sat_report_direct(submission_id)
            if 'error' in result:
                current_app.logger.error(f"Direct generation failed: {result['error']}")
                flash(f"Error generating report: {result['error']}", 'error')
                return redirect(url_for('dashboard.home'))
        else:
            current_app.logger.info(f"Template-based generation successful: {result['path']}")

        file_path = result['path']
        download_name = result['download_name']

        current_app.logger.info(f"Preparing download from path: {file_path}")
        
        # Read file and create response
        with open(file_path, 'rb') as f:
            file_content = f.read()
        
        response = Response(
            file_content,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename="{download_name}"',
                'Content-Length': str(len(file_content)),
                'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'Cache-Control': 'no-cache, no-store, must-revalidate',
                'Pragma': 'no-cache',
                'Expires': '0'
            }
        )
        
        current_app.logger.info(f"Serving SAT report download: {download_name} ({len(file_content)} bytes)")
        return response

    except Exception as e:
        current_app.logger.error(f"Error in download_report for {submission_id}: {e}", exc_info=True)
        flash('An unexpected error occurred while generating the report.', 'error')
        return redirect(url_for('dashboard.home'))


@status_bp.route('/download-modern/<submission_id>')
@login_required
def download_report_modern(submission_id):
    if not submission_id or submission_id == 'None':
        flash('Invalid submission ID.', 'error')
        return redirect(url_for('dashboard.home'))

    from services.report_renderer import generate_modern_sat_report

    result = generate_modern_sat_report(submission_id)
    if 'error' in result:
        flash(result['error'], 'error')
        return redirect(url_for('status.view_status', submission_id=submission_id))

    try:
        from services.file_download import safe_send_file
        
        response = safe_send_file(result['path'], result['download_name'], as_attachment=True)
        
        # Check if the response is an error (JSON response)
        if hasattr(response, 'status_code') and response.status_code != 200:
            flash('Error downloading modern report. The file may be corrupted.', 'error')
            return redirect(url_for('status.view_status', submission_id=submission_id))
        
        return response
    except Exception as exc:  # noqa: BLE001 - provide user-facing feedback
        current_app.logger.error(f'Error sending modern report for {submission_id}: {exc}', exc_info=True)
        flash('Error downloading report.', 'error')
        return redirect(url_for('status.view_status', submission_id=submission_id))


@status_bp.route('/diagnose/<submission_id>')
@login_required
def diagnose_download(submission_id):
    """Comprehensive diagnosis of download issues"""
    try:
        from models import Report
        import zipfile
        import hashlib
        
        # Only allow for admin users or in development
        if not (current_user.role == 'Admin' or current_app.config.get('DEBUG', False)):
            flash('Access denied.', 'error')
            return redirect(url_for('dashboard.home'))
        
        # Verify report exists
        report = Report.query.filter_by(id=submission_id).first()
        if not report:
            return jsonify({'error': 'Report not found'}), 404

        permanent_path = os.path.join(current_app.config['OUTPUT_DIR'], f'SAT_Report_{submission_id}_Final.docx')
        
        if not os.path.exists(permanent_path):
            return jsonify({'error': 'File not found', 'path': permanent_path}), 404
        
        diagnosis = {
            'file_path': permanent_path,
            'file_size': os.path.getsize(permanent_path),
            'tests': {}
        }
        
        # Test 1: File signature
        with open(permanent_path, 'rb') as f:
            first_bytes = f.read(50)
            diagnosis['tests']['file_signature'] = {
                'first_4_bytes': first_bytes[:4].hex(),
                'is_zip_signature': first_bytes.startswith(b'PK\x03\x04'),
                'first_50_bytes_hex': first_bytes.hex()
            }
        
        # Test 2: ZIP file validation
        try:
            with zipfile.ZipFile(permanent_path, 'r') as zip_file:
                file_list = zip_file.namelist()
                diagnosis['tests']['zip_validation'] = {
                    'is_valid_zip': True,
                    'file_count': len(file_list),
                    'has_word_directory': any(f.startswith('word/') for f in file_list),
                    'has_content_types': '[Content_Types].xml' in file_list,
                    'has_document_xml': 'word/document.xml' in file_list,
                    'sample_files': file_list[:10]
                }
        except Exception as e:
            diagnosis['tests']['zip_validation'] = {
                'is_valid_zip': False,
                'error': str(e)
            }
        
        # Test 3: File hash (to check for corruption)
        with open(permanent_path, 'rb') as f:
            file_hash = hashlib.md5(f.read()).hexdigest()
            diagnosis['tests']['file_integrity'] = {
                'md5_hash': file_hash,
                'file_readable': True
            }
        
        # Test 4: Try to read as DOCX with python-docx
        try:
            from docx import Document
            doc = Document(permanent_path)
            diagnosis['tests']['docx_library'] = {
                'can_open_with_python_docx': True,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
        except Exception as e:
            diagnosis['tests']['docx_library'] = {
                'can_open_with_python_docx': False,
                'error': str(e)
            }
        
        # Test 5: Content type detection
        import mimetypes
        detected_type, _ = mimetypes.guess_type(permanent_path)
        diagnosis['tests']['mime_detection'] = {
            'detected_type': detected_type,
            'expected_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        
        return jsonify(diagnosis)
        
    except Exception as e:
        current_app.logger.error(f"Error in diagnose_download for {submission_id}: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@status_bp.route('/test-docx/<submission_id>')
@login_required
def test_docx(submission_id):
    """Test endpoint to verify DOCX file can be opened"""
    try:
        from models import Report
        import zipfile
        
        # Only allow for admin users or in development
        if not (current_user.role == 'Admin' or current_app.config.get('DEBUG', False)):
            flash('Access denied.', 'error')
            return redirect(url_for('dashboard.home'))
        
        # Verify report exists
        report = Report.query.filter_by(id=submission_id).first()
        if not report:
            return jsonify({'error': 'Report not found'}), 404

        permanent_path = os.path.join(current_app.config['OUTPUT_DIR'], f'SAT_Report_{submission_id}_Final.docx')
        
        if not os.path.exists(permanent_path):
            return jsonify({'error': 'File not found', 'path': permanent_path}), 404
        
        # Test if it's a valid ZIP/DOCX file
        try:
            with zipfile.ZipFile(permanent_path, 'r') as zip_file:
                file_list = zip_file.namelist()
                has_word_files = any(f.startswith('word/') for f in file_list)
                has_content_types = '[Content_Types].xml' in file_list
                
                test_results = {
                    'file_path': permanent_path,
                    'file_size': os.path.getsize(permanent_path),
                    'is_valid_zip': True,
                    'file_count': len(file_list),
                    'has_word_files': has_word_files,
                    'has_content_types': has_content_types,
                    'is_valid_docx': has_word_files and has_content_types,
                    'sample_files': file_list[:10]  # First 10 files
                }
                
        except zipfile.BadZipFile:
            test_results = {
                'file_path': permanent_path,
                'file_size': os.path.getsize(permanent_path),
                'is_valid_zip': False,
                'error': 'File is not a valid ZIP/DOCX file'
            }
        
        return jsonify(test_results)
        
    except Exception as e:
        current_app.logger.error(f"Error in test_docx for {submission_id}: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@status_bp.route('/debug-file/<submission_id>')
@login_required
def debug_file(submission_id):
    """Debug endpoint to inspect file contents"""
    try:
        from models import Report, SATReport
        import json
        
        # Only allow for admin users or in development
        if not (current_user.role == 'Admin' or current_app.config.get('DEBUG', False)):
            flash('Access denied.', 'error')
            return redirect(url_for('dashboard.home'))
        
        # Verify report exists
        report = Report.query.filter_by(id=submission_id).first()
        if not report:
            return jsonify({'error': 'Report not found'}), 404

        permanent_path = os.path.join(current_app.config['OUTPUT_DIR'], f'SAT_Report_{submission_id}_Final.docx')
        
        if not os.path.exists(permanent_path):
            return jsonify({'error': 'File not found', 'path': permanent_path}), 404
        
        # Read first 2048 bytes to inspect content
        with open(permanent_path, 'rb') as f:
            first_bytes = f.read(2048)
            
        # Try to decode as text to see if it's HTML
        try:
            text_content = first_bytes.decode('utf-8', errors='ignore')
            is_text = True
        except:
            text_content = "Binary content"
            is_text = False
        
        file_info = {
            'file_path': permanent_path,
            'file_size': os.path.getsize(permanent_path),
            'first_4_bytes': first_bytes[:4],
            'first_4_bytes_hex': first_bytes[:4].hex(),
            'is_text_content': is_text,
            'first_200_chars': text_content[:200] if is_text else "Binary data",
            'contains_html': '<html' in text_content.lower() or '<!doctype' in text_content.lower() if is_text else False,
            'contains_pk_signature': first_bytes.startswith(b'PK\x03\x04'),
        }
        
        return jsonify(file_info)
        
    except Exception as e:
        current_app.logger.error(f"Error in debug_file for {submission_id}: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@status_bp.route('/simple-test/<submission_id>')
def simple_test(submission_id):
    """Simple test that doesn't require login - for debugging"""
    try:
        permanent_path = os.path.join(current_app.config['OUTPUT_DIR'], f'SAT_Report_{submission_id}_Final.docx')
        
        if not os.path.exists(permanent_path):
            return f"File not found at: {permanent_path}"
        
        # Simple file info
        file_size = os.path.getsize(permanent_path)
        
        # Check if it's a valid ZIP/DOCX
        with open(permanent_path, 'rb') as f:
            header = f.read(4)
        
        is_zip = header == b'PK\x03\x04'
        info = f"""
        <h2>File Information</h2>
        <p><strong>File Path:</strong> {permanent_path}</p>
        <p><strong>File Size:</strong> {file_size} bytes</p>
        <p><strong>Header:</strong> {header}</p>
        <p><strong>Is ZIP signature:</strong> {is_zip}</p>
        
        <h3>Download Links:</h3>
        <p><a href="/status/download/{submission_id}">New Direct Generator Download</a></p>
        <p><a href="/status/download-modern/{submission_id}">Modern Download (Old System)</a></p>
        """
        
        return info
        
    except Exception as e:
        return f"Error: {str(e)}"


@status_bp.route('/list')
@login_required
def list_submissions():
    """List all submissions for admin view"""
    from models import Report, SATReport

    try:
        # SQLAlchemy attaches columns dynamically; getattr avoids static analysis warnings.
        created_at_column = getattr(Report, 'created_at', None)
        query = Report.query
        if created_at_column is not None:
            query = query.order_by(created_at_column.desc())
        else:
            fallback_id = getattr(Report, 'id')
            query = query.order_by(fallback_id.desc())
        reports = query.all()
        submission_list = []

        for report in reports:
            sat_report = SATReport.query.filter_by(report_id=report.id).first()
            if not sat_report:
                continue

            try:
                stored_data = json.loads(sat_report.data_json)
            except json.JSONDecodeError:
                stored_data = {}

            # Determine overall status
            if report.approvals_json:
                try:
                    approvals = json.loads(report.approvals_json)
                    statuses = [a.get("status", "pending") for a in approvals]
                    if "rejected" in statuses:
                        overall_status = "rejected"
                    elif all(status == "approved" for status in statuses):
                        overall_status = "approved"
                    elif any(status == "approved" for status in statuses):
                        overall_status = "partially_approved"
                    else:
                        overall_status = "pending"
                except:
                    overall_status = "pending"
            else:
                overall_status = "draft"

            submission_list.append({
                "id": report.id,
                "document_title": stored_data.get("context", {}).get("DOCUMENT_TITLE", "SAT Report"),
                "client_name": stored_data.get("context", {}).get("CLIENT_NAME", ""),
                "created_at": report.created_at.strftime('%Y-%m-%d %H:%M:%S') if isinstance(report.created_at, dt.datetime) else report.created_at,
                "updated_at": report.updated_at.strftime('%Y-%m-%d %H:%M:%S') if isinstance(report.updated_at, dt.datetime) else report.updated_at,
                "status": overall_status,
                "user_email": report.user_email
            })

        return render_template('submissions_list.html', submissions=submission_list)

    except Exception as e:
        current_app.logger.error(f"Error fetching submissions list: {e}")
        flash('Error loading submissions.', 'error')
        return render_template('submissions_list.html', submissions=[])


