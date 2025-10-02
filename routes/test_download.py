from flask import Blueprint, send_file
from docx import Document
import os

test_download_bp = Blueprint('test_download', __name__)

@test_download_bp.route('/test_download')
def test_download():
    """Create and download a simple docx file"""
    document = Document()
    document.add_heading('Test Document', 0)
    document.add_paragraph('This is a test document.')
    
    file_path = os.path.join(os.getcwd(), 'test.docx')
    document.save(file_path)
    
    return send_file(file_path, as_attachment=True)
